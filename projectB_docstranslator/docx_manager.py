from pathlib import Path

from docx import Document

#2 read_docx_files 함수
def read_docx_files(file_path):
    doc = Document(file_path)
    doc_text = ''
    for para in doc.paragraphs:
        doc_text += para.text + '\n'

    # print(doc_text)
    return doc_text

#4 write_docx_files 함수
def write_docx_files(translated_data, output_directory, file_name):
    output_directory = Path(output_directory)
    output_directory.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_paragraph(translated_data)
    doc.save(output_directory / f'translated_{file_name}')
